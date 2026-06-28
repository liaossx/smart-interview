import fitz  # PyMuPDF
import docx


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文件，保留表格结构"""
    doc = fitz.open(file_path)
    parts = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        page_text = ""
        for block in blocks:
            if block["type"] == 0:  # 文本块
                for line in block["lines"]:
                    spans = line["spans"]
                    if spans:
                        # 同一行内的文本拼接
                        line_text = " ".join(s["text"] for s in spans).strip()
                        if line_text:
                            page_text += line_text + "\n"
            elif block["type"] == 1:  # 图片块
                page_text += "[图片]\n"
        if page_text.strip():
            parts.append(page_text.strip())

        # 尝试提取表格
        table_text = _extract_tables(page)
        if table_text:
            parts.append(table_text)

    doc.close()
    return "\n\n".join(parts).strip()


def _extract_tables(page) -> str:
    """提取页面内的表格为结构化文本"""
    try:
        tables = page.find_tables()
        if not tables or not tables.tables:
            return ""
        result = []
        for table in tables.tables:
            result.append("--- 表格 ---")
            header = table.header
            if header and header.names:
                result.append(" | ".join(str(n) for n in header.names if n))
                result.append("-" * 40)
            for row in table.extract():
                result.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
        return "\n".join(result)
    except Exception:
        return ""


def parse_docx(file_path: str) -> str:
    """解析 DOCX 文件，含表格提取"""
    doc = docx.Document(file_path)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            text = _get_paragraph_text(element, doc)
            if text.strip():
                parts.append(text.strip())
        elif tag == "tbl":
            table_text = _parse_docx_table(element, doc)
            if table_text:
                parts.append(table_text)

    return "\n\n".join(parts).strip() or "\n".join([p.text for p in doc.paragraphs]).strip()


def _get_paragraph_text(element, doc) -> str:
    """从 XML 段落元素提取文本"""
    texts = []
    for node in element.iter():
        tag = node.tag.split("}")[-1] if "}" in node.tag else node.tag
        if tag == "t" and node.text:
            texts.append(node.text)
    return "".join(texts)


def _parse_docx_table(element, doc) -> str:
    """解析 DOCX 表格为结构化文本"""
    rows = []
    for row_elem in element.iter():
        tag = row_elem.tag.split("}")[-1] if "}" in row_elem.tag else row_elem.tag
        if tag != "tr":
            continue
        cells = []
        for cell in row_elem.iter():
            ct = cell.tag.split("}")[-1] if "}" in cell.tag else cell.tag
            if ct != "tc":
                continue
            cell_text = ""
            for p in cell.iter():
                pt = p.tag.split("}")[-1] if "}" in p.tag else p.tag
                if pt == "t" and p.text:
                    cell_text += p.text
                elif pt == "p":
                    for t in p.iter():
                        tt = t.tag.split("}")[-1] if "}" in t.tag else t.tag
                        if tt == "t" and t.text:
                            cell_text += t.text
            cells.append(cell_text.strip())
        if cells:
            rows.append(" | ".join(cells))
    if rows:
        return "--- 表格 ---\n" + "\n".join(rows)
    return ""


def parse_document(file_path: str) -> str:
    """根据文件扩展名自动选择解析方式"""
    if file_path.endswith(".pdf"):
        return parse_pdf(file_path)
    elif file_path.endswith(".docx"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_path}")
